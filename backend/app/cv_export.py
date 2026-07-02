"""Render a tailored CV (markdown) into .docx or .pdf that matches the
candidate's master CV layout exactly.

Master CV spec (extracted from Saw_Thu_Naing_CV_Tamis.docx):
  Page:      A4, 0.75" margins all sides
  Font:      Arial throughout
  Name:      20pt bold
  Title:     11pt
  Contact:   9.5pt, pipe-separated
  Section:   11.5pt bold CAPS
  Entry:     10.5pt bold, role/company on left, dates RIGHT-aligned at 6.27"
  Bullet:    10pt, optional bold lead phrase

The tailored body comes from the AI as markdown. Experience/education headings
encode dates after " @@ " so we can right-align them.

Pure-Python: python-docx + reportlab. No system dependencies.
"""
from __future__ import annotations
import io
import re
from typing import Any

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _split_bold(text: str) -> list[tuple[str, bool]]:
    runs: list[tuple[str, bool]] = []
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False))
        runs.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False))
    return runs or [("", False)]


_DATE_TAIL_RE = re.compile(
    r"\s*[\(\-–—|]*\s*"
    r"((?:[A-Z][a-z]{2,9}\.?\s+)?\d{4}\s*[–-]\s*(?:Present|Current|(?:[A-Z][a-z]{2,9}\.?\s+)?\d{4}))"
    r"\s*\)?\s*$"
)


def _split_heading_dates(text: str) -> tuple[str, str | None]:
    if "@@" in text:
        left, _, right = text.partition("@@")
        return left.strip(" •-—|"), right.strip()
    m = _DATE_TAIL_RE.search(text)
    if m:
        return text[: m.start()].strip(" •-—|()"), m.group(1).strip()
    return text.strip(), None


def _parse_blocks(md: str) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    bullets: list[str] = []

    def flush():
        nonlocal bullets
        if bullets:
            blocks.append({"kind": "bullets", "items": bullets})
            bullets = []

    for raw in md.replace("\r\n", "\n").split("\n"):
        s = raw.strip()
        if not s:
            flush(); continue
        if s.startswith("### "):
            flush()
            left, dates = _split_heading_dates(s[4:].strip())
            blocks.append({"kind": "entry", "text": left, "dates": dates})
        elif s.startswith("## "):
            flush(); blocks.append({"kind": "h2", "text": s[3:].strip()})
        elif s.startswith("# "):
            flush(); blocks.append({"kind": "h1", "text": s[2:].strip()})
        elif s.startswith(("- ", "* ")):
            bullets.append(s[2:].strip())
        elif re.match(r"^\d+\.\s", s):
            bullets.append(re.sub(r"^\d+\.\s", "", s))
        else:
            flush(); blocks.append({"kind": "para", "text": s})
    flush()
    return blocks


def _contact_line(p) -> str:
    bits = [getattr(p, "location", None), getattr(p, "phone", None),
            getattr(p, "email", None), getattr(p, "website", None),
            getattr(p, "linkedin", None), getattr(p, "github", None)]
    return "  |  ".join(b for b in bits if b)


FONT = "Arial"
RIGHT_TAB_IN = 6.27  # matches master CV


# ---- DOCX -----------------------------------------------------------------

def build_docx(profile, content_md: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_TAB_ALIGNMENT

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.75)
    sec.left_margin = sec.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)

    BLACK = RGBColor(0x00, 0x00, 0x00)

    def add_runs(p, text, *, bold=False, size=10.0, color=BLACK):
        for seg, sb in _split_bold(text):
            r = p.add_run(seg)
            r.font.name = FONT
            r.bold = bold or sb
            r.font.size = Pt(size)
            r.font.color.rgb = color

    def spacing(p, before=0, after=2):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.0

    # Header
    name = getattr(profile, "full_name", None) or "Curriculum Vitae"
    p = doc.add_paragraph(); add_runs(p, name, bold=True, size=20); spacing(p, 0, 1)
    if getattr(profile, "headline", None):
        p = doc.add_paragraph(); add_runs(p, profile.headline, size=11); spacing(p, 0, 1)
    contact = _contact_line(profile)
    if contact:
        p = doc.add_paragraph(); add_runs(p, contact, size=9.5); spacing(p, 0, 8)

    blocks = _parse_blocks(content_md)
    if blocks and blocks[0]["kind"] == "h1" and blocks[0]["text"].lower() in (name.lower(), ""):
        blocks = blocks[1:]

    for b in blocks:
        if b["kind"] in ("h1", "h2"):
            p = doc.add_paragraph(); add_runs(p, b["text"].upper(), bold=True, size=11.5)
            spacing(p, 8, 3)
        elif b["kind"] == "entry":
            p = doc.add_paragraph()
            if b.get("dates"):
                p.paragraph_format.tab_stops.add_tab_stop(Inches(RIGHT_TAB_IN), WD_TAB_ALIGNMENT.RIGHT)
            add_runs(p, b["text"], bold=True, size=10.5)
            if b.get("dates"):
                r = p.add_run("\t"); r.font.name = FONT
                r2 = p.add_run(b["dates"]); r2.font.name = FONT; r2.bold = True; r2.font.size = Pt(10.5)
            spacing(p, 6, 1)
        elif b["kind"] == "bullets":
            for it in b["items"]:
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, it, size=10)
                spacing(p, 0, 1)
        else:
            p = doc.add_paragraph(); add_runs(p, b["text"], size=10); spacing(p, 0, 3)

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ---- PDF ------------------------------------------------------------------

def build_pdf(profile, content_md: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch, mm
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.lib.colors import black
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, ListFlowable, ListItem, Table, TableStyle,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    ss = getSampleStyleSheet()
    F, FB = "Helvetica", "Helvetica-Bold"  # Helvetica == Arial metrics in PDF

    name_s = ParagraphStyle("N", parent=ss["Normal"], fontName=FB, fontSize=20,
                            textColor=black, leading=23, spaceAfter=1)
    title_s = ParagraphStyle("T", parent=ss["Normal"], fontName=F, fontSize=11,
                             textColor=black, leading=14, spaceAfter=1)
    contact_s = ParagraphStyle("C", parent=ss["Normal"], fontName=F, fontSize=9.5,
                               textColor=black, leading=12, spaceAfter=8)
    h2_s = ParagraphStyle("H2", parent=ss["Normal"], fontName=FB, fontSize=11.5,
                          textColor=black, spaceBefore=9, spaceAfter=3, leading=14)
    el_s = ParagraphStyle("EL", parent=ss["Normal"], fontName=FB, fontSize=10.5,
                          textColor=black, leading=13, alignment=TA_LEFT)
    er_s = ParagraphStyle("ER", parent=ss["Normal"], fontName=FB, fontSize=10.5,
                          textColor=black, leading=13, alignment=TA_RIGHT)
    body_s = ParagraphStyle("B", parent=ss["Normal"], fontName=F, fontSize=10,
                            textColor=black, leading=13, alignment=TA_LEFT, spaceAfter=2)
    bullet_s = ParagraphStyle("BU", parent=body_s, spaceAfter=1.5)

    def rl(text: str) -> str:
        out = []
        for seg, b in _split_bold(text):
            seg = seg.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            out.append(f"<b>{seg}</b>" if b else seg)
        return "".join(out)

    page_w, _ = A4
    margin = 0.75 * inch
    content_w = page_w - 2 * margin

    story = []
    name = getattr(profile, "full_name", None) or "Curriculum Vitae"
    story.append(Paragraph(rl(name), name_s))
    if getattr(profile, "headline", None):
        story.append(Paragraph(rl(profile.headline), title_s))
    contact = _contact_line(profile)
    if contact:
        story.append(Paragraph(rl(contact), contact_s))

    blocks = _parse_blocks(content_md)
    if blocks and blocks[0]["kind"] == "h1" and blocks[0]["text"].lower() in (name.lower(), ""):
        blocks = blocks[1:]

    for b in blocks:
        if b["kind"] in ("h1", "h2"):
            story.append(Paragraph(rl(b["text"].upper()), h2_s))
        elif b["kind"] == "entry":
            if b.get("dates"):
                tbl = Table([[Paragraph(rl(b["text"]), el_s), Paragraph(rl(b["dates"]), er_s)]],
                            colWidths=[content_w * 0.74, content_w * 0.26])
                tbl.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ]))
                story.append(tbl)
            else:
                story.append(Paragraph(rl(b["text"]), el_s))
        elif b["kind"] == "bullets":
            items = [ListItem(Paragraph(rl(it), bullet_s), leftIndent=9) for it in b["items"]]
            story.append(ListFlowable(items, bulletType="bullet", start="•",
                                      bulletColor=black, leftIndent=11, bulletFontSize=7,
                                      spaceBefore=2, spaceAfter=2))
        else:
            story.append(Paragraph(rl(b["text"]), body_s))

    buf = io.BytesIO()
    SimpleDocTemplate(buf, pagesize=A4, topMargin=margin, bottomMargin=margin,
                      leftMargin=margin, rightMargin=margin, title="Tailored CV").build(story)
    return buf.getvalue()
