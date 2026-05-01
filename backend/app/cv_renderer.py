"""Render a tailored CV as PDF or DOCX.

Strategy:
- Identity (name, contact) comes from Profile (DB).
- Body content (summary, experience, skills, etc.) comes from the AI's
  tailored markdown stored in job.tailored_docs['cv']['content'].
- The same parsed structure feeds both PDF (via Jinja+WeasyPrint) and DOCX
  (via python-docx) so output stays consistent across formats.
"""
from __future__ import annotations
import io
import re
from pathlib import Path
from typing import Optional

import markdown2
from jinja2 import Environment, FileSystemLoader, select_autoescape
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML

from . import models


# ---------------------------------------------------------------------------
# Markdown parsing - turn the AI's section-y markdown into a small structured
# dict we can feed to either renderer.
# ---------------------------------------------------------------------------

# The AI is instructed to produce H2 sections like "## Summary", "## Experience".
# We split on those markers.
_SECTION_RE = re.compile(r"^##\s+(.+?)\s*$", re.MULTILINE)


def _parse_sections(md: str) -> dict[str, str]:
    """Split markdown into a dict of {lowercased_title: markdown_body}."""
    if not md:
        return {}
    # Find all section headers and their positions
    matches = list(_SECTION_RE.finditer(md))
    if not matches:
        # No section headers - treat the whole thing as a generic body
        return {"_body": md.strip()}
    sections: dict[str, str] = {}
    for i, m in enumerate(matches):
        title = m.group(1).strip().lower()
        body_start = m.end()
        body_end = matches[i + 1].start() if i + 1 < len(matches) else len(md)
        body = md[body_start:body_end].strip()
        sections[title] = body
    return sections


def _md_to_html(md: str) -> str:
    """Convert markdown body to HTML using sensible options."""
    if not md:
        return ""
    return markdown2.markdown(md, extras=["fenced-code-blocks", "cuddled-lists"])


# ---------------------------------------------------------------------------
# Match section names robustly. AI output varies ("Summary", "Professional Summary",
# "Profile" all mean the same thing).
# ---------------------------------------------------------------------------
_SECTION_ALIASES = {
    "summary":      ["summary", "professional summary", "profile", "about"],
    "experience":   ["experience", "professional experience", "work experience", "employment"],
    "education":    ["education", "academic background"],
    "skills":       ["skills", "technical skills", "core skills", "key skills"],
    "achievements": ["awards", "awards & achievements", "achievements", "certifications", "awards and achievements"],
}


def _extract(sections: dict[str, str], key: str) -> str:
    """Find the section body matching any alias for `key`."""
    for alias in _SECTION_ALIASES[key]:
        if alias in sections:
            return sections[alias]
    return ""


# ---------------------------------------------------------------------------
# Template loading
# ---------------------------------------------------------------------------
_TEMPLATE_DIR = Path(__file__).parent / "templates"
_jinja_env = Environment(
    loader=FileSystemLoader(str(_TEMPLATE_DIR)),
    autoescape=select_autoescape(["html"]),
)


# ---------------------------------------------------------------------------
# PDF rendering
# ---------------------------------------------------------------------------

def _build_summary_text(sections: dict[str, str]) -> str:
    """The summary section is plain prose, not bullets - return as plain text."""
    raw = _extract(sections, "summary")
    if not raw:
        return ""
    # Strip stray markdown emphasis but keep newlines collapsed
    text = re.sub(r"[*_`]+", "", raw)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def render_pdf(profile: Optional[models.Profile], tailored_md: str) -> bytes:
    """Render a tailored CV as a PDF binary.

    Profile supplies identity (name, contact) - never trust the AI to output
    these. tailored_md is the AI-generated body markdown.
    """
    sections = _parse_sections(tailored_md)

    # Build identity from profile (with sane fallbacks)
    name = (profile.full_name if profile and profile.full_name else "Candidate") or "Candidate"
    email = profile.email if profile else None
    phone = profile.phone if profile else None
    linkedin = profile.linkedin if profile else None
    github = profile.github if profile else None

    template = _jinja_env.get_template("cv.html")
    html_str = template.render(
        name=name,
        email=email,
        phone=phone,
        linkedin=linkedin,
        github=github,
        summary=_build_summary_text(sections),
        experience_html=_md_to_html(_extract(sections, "experience")),
        education_html=_md_to_html(_extract(sections, "education")),
        skills_html=_md_to_html(_extract(sections, "skills")),
        achievements_html=_md_to_html(_extract(sections, "achievements")),
    )

    pdf_bytes = HTML(string=html_str).write_pdf()
    return pdf_bytes


# ---------------------------------------------------------------------------
# DOCX rendering - parallel structure, ATS-safe styling
# ---------------------------------------------------------------------------

def _set_run_font(run, size_pt: float = 10.5, bold: bool = False, italic: bool = False):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_run_font(run, size_pt=11, bold=True)
    # Add bottom border by inserting a paragraph property
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_markdown_body(doc: Document, md: str):
    """Naive markdown-to-docx: handles bullets, sub-headings, and paragraphs.
    ATS-safe: standard bullets, no fancy formatting.
    """
    if not md:
        return
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        # H3 sub-heading (e.g. job title line)
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line[4:].strip())
            _set_run_font(run, size_pt=10.5, bold=True)
            i += 1
            continue

        # Bullet
        if line.lstrip().startswith(("- ", "* ", "• ")):
            text = line.lstrip()[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                _set_run_font(run, size_pt=10.5)
            run = p.add_run(text) if not p.runs else None
            if run is not None:
                _set_run_font(run, size_pt=10.5)
            else:
                # paragraph already has style+text - replace its content
                p.clear()  # not actually needed; List Bullet style created with text via add_run
                run = p.add_run(text)
                _set_run_font(run, size_pt=10.5)
            i += 1
            continue

        # Plain paragraph - clean stray markdown emphasis chars
        text = re.sub(r"[*_`]+", "", line).strip()
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            _set_run_font(run, size_pt=10.5)
        i += 1


def render_docx(profile: Optional[models.Profile], tailored_md: str) -> bytes:
    """Render a tailored CV as a DOCX binary."""
    sections = _parse_sections(tailored_md)

    doc = Document()
    # Tighten margins
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Identity
    name = (profile.full_name if profile and profile.full_name else "Candidate") or "Candidate"

    # Name centered, large, bold
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name.upper())
    _set_run_font(run, size_pt=18, bold=True)

    # Contact line
    contact_parts = []
    if profile:
        for v in (profile.email, profile.phone, profile.linkedin, profile.github):
            if v:
                contact_parts.append(v)
    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(" | ".join(contact_parts))
        _set_run_font(run, size_pt=10)

    # Sections in fixed order
    summary = _build_summary_text(sections)
    if summary:
        _add_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(summary)
        _set_run_font(run, size_pt=10.5)

    if _extract(sections, "experience"):
        _add_section_heading(doc, "Professional Experience")
        _add_markdown_body(doc, _extract(sections, "experience"))

    if _extract(sections, "education"):
        _add_section_heading(doc, "Education")
        _add_markdown_body(doc, _extract(sections, "education"))

    if _extract(sections, "skills"):
        _add_section_heading(doc, "Technical Skills")
        _add_markdown_body(doc, _extract(sections, "skills"))

    if _extract(sections, "achievements"):
        _add_section_heading(doc, "Awards & Achievements")
        _add_markdown_body(doc, _extract(sections, "achievements"))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Filename generation
# ---------------------------------------------------------------------------
def safe_filename(name: str) -> str:
    """Sanitize a string for use as a filename."""
    s = re.sub(r"[^A-Za-z0-9_\- ]+", "", name).strip()
    s = re.sub(r"\s+", "_", s)
    return s or "Untitled"


def build_filename(profile: Optional[models.Profile], company: str, role: str, ext: str) -> str:
    name = profile.full_name if profile and profile.full_name else "CV"
    return f"{safe_filename(name)}_{safe_filename(role)}_{safe_filename(company)}.{ext}"
